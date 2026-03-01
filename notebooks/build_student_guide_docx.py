"""
build_student_guide_docx.py
Converts Presentation_Guide_Student.md to a styled Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re

# ── Paths ──────────────────────────────────────────────────────────────────
BASE    = Path('/Volumes/D Drive/Data analysis/Olympic data analysis')
SRC     = BASE / 'outputs' / 'report' / 'Presentation_Guide_Student.md'
OUT     = BASE / 'outputs' / 'presentation' / 'Presentation_Guide_Student.docx'

# ── Colors ─────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x00, 0x85, 0xC7)
YELLOW = RGBColor(0xF4, 0xC3, 0x00)
GREEN  = RGBColor(0x00, 0x9F, 0x6B)
RED    = RGBColor(0xDF, 0x00, 0x24)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREY   = RGBColor(0x44, 0x44, 0x44)
LGREY  = RGBColor(0xDD, 0xDD, 0xDD)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Document setup ─────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Helper: shade table cell ───────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.lstrip('#'))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: add styled run within a paragraph ─────────────────────────────
def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return run

# ── Helper: parse inline markdown (bold/italic) ───────────────────────────
def add_inline(para, text, base_bold=False, base_color=None, base_size=None):
    """Parse **bold**, *italic*, and plain text, adding runs to para."""
    # Pattern: **bold**, *italic*, or plain text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    for match in re.finditer(pattern, text):
        chunk = match.group()
        if chunk.startswith('**') and chunk.endswith('**'):
            add_run(para, chunk[2:-2], bold=True, color=base_color, size=base_size)
        elif chunk.startswith('*') and chunk.endswith('*'):
            add_run(para, chunk[1:-1], italic=True, color=base_color, size=base_size)
        else:
            add_run(para, chunk, bold=base_bold, color=base_color, size=base_size)

# ── Helper: horizontal rule ────────────────────────────────────────────────
def add_hr(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
# Olympic colour band (table with 5 cells)
tbl = doc.add_table(rows=1, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
colors = ['0085C7', 'F4C300', '009F6B', 'DF0024', '000000']
for i, cell in enumerate(tbl.rows[0].cells):
    shade_cell(cell, colors[i])
    cell.width = Inches(1.3)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)

doc.add_paragraph()  # spacer

# Title
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('LA 2028 Olympic Games')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_para.add_run('Presentation Guide')
run2.font.size = Pt(22)
run2.font.bold = True
run2.font.color.rgb = BLACK
run2.font.name = 'Calibri'

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2.add_run('Written for a 10th Grade Student')
run3.font.size = Pt(14)
run3.font.italic = True
run3.font.color.rgb = GREY
run3.font.name = 'Calibri'

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = meta.add_run('SportsFanatics Consulting Agency  |  March 2026  |  Analyst: Shabeeb')
run4.font.size = Pt(10)
run4.font.color.rgb = GREY
run4.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PARSE AND RENDER MARKDOWN
# ══════════════════════════════════════════════════════════════════════════
md_text = SRC.read_text(encoding='utf-8')
lines   = md_text.split('\n')

# Slide heading color rotation (matching Olympic rings)
SLIDE_COLORS = [BLUE, RED, GREEN, BLUE, RED, GREEN, BLUE, RED, GREEN,
                BLUE, RED, GREEN, BLUE, RED, GREEN, BLUE, RED, GREEN,
                BLUE, RED, GREEN]
slide_idx = 0

# Track state
in_table      = False
table_rows    = []
in_blockquote = False
in_numbered   = False
numbered_items= []
in_bulleted   = False
bullet_items  = []

def flush_table(doc, rows):
    """Render a markdown table into a Word table."""
    if not rows:
        return
    # rows[0] = header, rows[1] = separator, rows[2:] = data
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    data   = [[c.strip() for c in r.strip('|').split('|')] for r in rows[2:]]

    tbl = doc.add_table(rows=1 + len(data), cols=len(header))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = tbl.rows[0]
    for i, hdr in enumerate(header):
        cell = hrow.cells[i]
        shade_cell(cell, '0085C7')
        para = cell.paragraphs[0]
        run  = para.add_run(hdr)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(10)
        run.font.name  = 'Calibri'

    # Data rows
    for r_i, row in enumerate(data):
        drow = tbl.rows[r_i + 1]
        fill = 'F0F4F8' if r_i % 2 == 0 else 'FFFFFF'
        for c_i, val in enumerate(row):
            cell = drow.cells[c_i]
            shade_cell(cell, fill)
            para = cell.paragraphs[0]
            # Check for **bold** in cell
            add_inline(para, val, base_size=10)

    doc.add_paragraph()

def flush_bullets(doc, items, numbered=False):
    if not items:
        return
    for i, item in enumerate(items):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Inches(0.3)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        prefix = f'{i+1}. ' if numbered else '• '
        add_run(para, prefix, bold=False, color=BLUE)
        add_inline(para, item)

i = 0
while i < len(lines):
    line = lines[i]

    # ── Collect table block ────────────────────────────────────────────
    if line.startswith('|'):
        if not in_table:
            in_table   = True
            table_rows = []
        table_rows.append(line)
        i += 1
        continue
    elif in_table:
        flush_table(doc, table_rows)
        in_table   = False
        table_rows = []

    # ── Flush pending bullets / numbered ──────────────────────────────
    stripped = line.strip()

    if not (stripped.startswith('- ') or stripped.startswith('* ')) and bullet_items:
        flush_bullets(doc, bullet_items, numbered=False)
        bullet_items = []

    num_match = re.match(r'^\d+\. (.+)', stripped)
    if not num_match and numbered_items:
        flush_bullets(doc, numbered_items, numbered=True)
        numbered_items = []

    # ── Horizontal rule ────────────────────────────────────────────────
    if stripped in ('---', '***', '___'):
        add_hr(doc)
        i += 1
        continue

    # ── H1 ─────────────────────────────────────────────────────────────
    if line.startswith('# '):
        text = line[2:].strip()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(4)
        run = para.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = BLUE
        run.font.name = 'Calibri'
        i += 1
        continue

    # ── H2 ─────────────────────────────────────────────────────────────
    if line.startswith('## '):
        text = line[3:].strip()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(4)

        # Detect "Slide N" headings
        is_slide = text.lower().startswith('slide ')
        color = SLIDE_COLORS[slide_idx % len(SLIDE_COLORS)] if is_slide else BLUE
        if is_slide:
            slide_idx += 1

        run = para.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = color
        run.font.name = 'Calibri'

        # Underline for slide headings
        if is_slide:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*color))
            pBdr.append(bottom)
            pPr.append(pBdr)

        i += 1
        continue

    # ── H3 ─────────────────────────────────────────────────────────────
    if line.startswith('### '):
        text = line[4:].strip()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = GREEN
        run.font.name = 'Calibri'
        i += 1
        continue

    # ── Blockquote ─────────────────────────────────────────────────────
    if line.startswith('> '):
        text = line[2:].strip()
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Inches(0.4)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '0085C7')
        pBdr.append(left)
        pPr.append(pBdr)
        add_inline(para, text, base_color=GREY)
        i += 1
        continue

    # ── Numbered list ──────────────────────────────────────────────────
    if num_match:
        numbered_items.append(num_match.group(1))
        i += 1
        continue

    # ── Bullet list ────────────────────────────────────────────────────
    if stripped.startswith('- ') or stripped.startswith('* '):
        item = stripped[2:].strip()
        bullet_items.append(item)
        i += 1
        continue

    # ── Blank line ─────────────────────────────────────────────────────
    if stripped == '':
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        i += 1
        continue

    # ── Italic-only line (footer lines starting with *...*) ────────────
    if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
        text = stripped[1:-1]
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        run = para.add_run(text)
        run.font.italic = True
        run.font.size   = Pt(10)
        run.font.color.rgb = GREY
        run.font.name   = 'Calibri'
        i += 1
        continue

    # ── Regular paragraph ──────────────────────────────────────────────
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    add_inline(para, stripped)
    i += 1

# Flush any remaining lists/tables
if bullet_items:
    flush_bullets(doc, bullet_items, numbered=False)
if numbered_items:
    flush_bullets(doc, numbered_items, numbered=True)
if in_table and table_rows:
    flush_table(doc, table_rows)

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1024
print(f'Saved: {OUT}')
print(f'Size:  {size_kb:.0f} KB')
print('Done.')
